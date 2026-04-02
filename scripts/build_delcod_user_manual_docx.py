from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Projeto Leitor de cod de barras\.worktrees\v1-bootstrap")
TARGET = ROOT / "docs" / "manual" / "2026-04-02-delcod-manual-de-uso.docx"
DESKTOP_IMAGE = ROOT / "docs" / "manual" / "assets" / "delcod-home.png"
MOBILE_IMAGE = ROOT / "docs" / "manual" / "assets" / "delcod-mobile-home.png"


def set_run_font(run, *, name: str = "Aptos", size: int = 11, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
    set_run_font(run, size=10)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)


def ensure_style(document: Document, name: str):
    if name in document.styles:
        return document.styles[name]
    return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(14, 72, 68)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Aptos"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(17, 24, 39)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)

    title = ensure_style(document, "ManualTitle")
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(14, 72, 68)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    subtitle = ensure_style(document, "ManualSubtitle")
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    subtitle.font.size = Pt(12)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)

    caption = ensure_style(document, "ManualCaption")
    caption.font.name = "Aptos"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)

    note = ensure_style(document, "ManualNote")
    note.font.name = "Aptos"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    note.font.size = Pt(10)
    note.paragraph_format.space_after = Pt(6)


def add_cover(document: Document):
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph(style="ManualTitle")
    p.add_run("DelCod")

    p = document.add_paragraph(style="ManualSubtitle")
    p.add_run("Manual de Uso Detalhado")

    p = document.add_paragraph(style="ManualSubtitle")
    p.add_run("Versão operacional do aplicativo para inventário de bobinas")

    for _ in range(12):
        document.add_paragraph("")

    p = document.add_paragraph(style="ManualSubtitle")
    p.add_run("Atualizado em 02/04/2026")

    p = document.add_paragraph(style="ManualSubtitle")
    p.add_run("Documento interno de apoio ao uso do sistema")


def add_paragraph(document: Document, text: str, *, style: str = "Normal"):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_bullets(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        run = p.add_run(f"• {item}")
        set_run_font(run)


def add_numbered_steps(document: Document, steps: list[str]):
    for index, step in enumerate(steps, start=1):
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(f"{index}. {step}")
        set_run_font(run)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float):
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        add_paragraph(document, caption, style="ManualCaption")


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        set_run_font(run, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(value)
            set_run_font(run)


def build_document() -> Document:
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    add_cover(document)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    add_page_number(body_section.footer.paragraphs[0])

    document.add_heading("1. Objetivo do DelCod", level=1)
    add_paragraph(
        document,
        "O DelCod é um aplicativo de inventário de bobinas criado para registrar, consultar, importar, exportar e manter sincronizada uma lista global de lotes de bobina. O uso diário do sistema foi pensado para duas situações principais: leitura por câmera no Android e entrada manual no navegador.",
    )
    add_bullets(
        document,
        [
            "Registrar lotes de bobina rapidamente no galpão ou no escritório.",
            "Associar cada lote a um armazém e, por consequência, à empresa correta.",
            "Importar bases existentes em Excel ou CSV para acelerar o inventário.",
            "Exportar o resultado em Excel ou PDF para conferência, análise e envio.",
            "Continuar operando mesmo sem internet, com sincronização automática quando a conexão volta.",
        ],
    )

    document.add_heading("2. Formas de acesso", level=1)
    add_bullets(
        document,
        [
            "Navegador: acesso pelo link oficial do sistema, indicado para entrada manual, conferência, importação e exportação.",
            "Android: uso do APK interno do DelCod, indicado para leitura por câmera e inventário no chão de fábrica ou galpão.",
            "Sincronização: sempre que houver internet, os dispositivos atualizam a mesma lista global em tempo real.",
        ],
    )
    add_paragraph(
        document,
        "Endereço web atual: https://delcod-phi.vercel.app",
        style="ManualNote",
    )

    document.add_heading("3. Visão geral da tela principal", level=1)
    add_paragraph(
        document,
        "A tela principal do DelCod foi desenhada para concentrar todas as funções do processo em um único lugar. No navegador, o painel de entrada e ações fica à esquerda e a lista global à direita. No Android, a tela inteira rola em um fluxo único, com os blocos subindo juntos conforme o usuário navega.",
    )
    add_image(
        document,
        DESKTOP_IMAGE,
        "Figura 1. Tela principal do DelCod no navegador, com painel de entrada, indicadores, ações da lista e lista global.",
        16.5,
    )
    add_image(
        document,
        MOBILE_IMAGE,
        "Figura 2. Exemplo do DelCod em largura móvel, mostrando a organização compacta para uso no celular.",
        7.2,
    )
    add_bullets(
        document,
        [
            "Cabeçalho com o nome DelCod e os números de apoio ao trabalho diário.",
            "Banner de sincronização com o estado atual da comunicação com o banco de dados.",
            "Área principal de leitura ou entrada manual do lote de bobina.",
            "Resumo rápido com os indicadores Ativos, Selecionados e Pendentes.",
            "Ações da lista para importar, exportar, selecionar, limpar e alocar armazém em lote.",
            "Lista global com todos os lotes registrados em tempo real.",
        ],
    )

    document.add_heading("4. Indicadores e status", level=1)
    add_paragraph(
        document,
        "Os indicadores do topo ajudam a interpretar rapidamente a situação do inventário e o estado da sincronização.",
    )
    add_simple_table(
        document,
        ["Indicador ou status", "Significado operacional"],
        [
            ["Ativos", "Quantidade total de lotes atualmente registrados na lista global."],
            ["Selecionados", "Quantidade de itens marcados para ação em lote ou exportação."],
            ["Pendentes", "Quantidade de lotes que ainda estão sem armazém alocado."],
            ["Sincronizado", "O app está online e sem pendências de envio."],
            ["Sincronizando", "Existem operações sendo enviadas para o banco."],
            ["Offline", "O dispositivo está sem conexão; o app continua funcionando localmente."],
            ["Falha na sincronização", "Houve erro no envio de dados e o sistema tentará novamente."],
        ],
    )
    add_paragraph(
        document,
        "Quando um lote aparece com o status 'Sem armazém alocado', ele não está perdido. Isso significa apenas que já foi registrado, mas ainda precisa de definição de armazém.",
    )

    document.add_heading("5. Registrar lotes de bobina", level=1)
    document.add_heading("5.1 Leitura por câmera no Android", level=2)
    add_paragraph(
        document,
        "No Android, a leitura por câmera é a forma mais rápida de alimentar o inventário. O scanner foi ajustado para priorizar leitura estável e rápida de códigos lineares, com tolerância maior a pequenas falhas físicas na etiqueta.",
    )
    add_numbered_steps(
        document,
        [
            "Abra o DelCod no Android.",
            "Na área 'Leitura por câmera', aponte a câmera para a etiqueta da bobina.",
            "Quando o código for reconhecido, o lote será registrado no sistema.",
            "Se o armazém já estiver selecionado no painel, ele será aplicado imediatamente.",
            "Se nenhum armazém estiver definido, o lote será salvo como pendente para ajuste posterior.",
        ],
    )

    document.add_heading("5.2 Entrada manual no navegador", level=2)
    add_paragraph(
        document,
        "No navegador, o fluxo principal é a digitação ou colagem manual do lote de bobina. Esse modo é ideal para conferência, ajustes pontuais e lançamentos administrativos.",
    )
    add_numbered_steps(
        document,
        [
            "Digite ou cole o valor no campo 'Lote de Bobina'.",
            "Se já souber o armazém, selecione-o no campo 'Armazém'.",
            "Clique em 'Adicionar'.",
            "O lote entra na lista global imediatamente.",
        ],
    )

    document.add_heading("5.3 Armazém e empresa derivada", level=2)
    add_paragraph(
        document,
        "O DelCod usa o armazém para identificar automaticamente a empresa responsável pela bobina. Essa derivação é automática e não precisa ser informada manualmente.",
    )
    add_simple_table(
        document,
        ["Armazém", "Empresa derivada automaticamente"],
        [
            ["05", "Bora Embalagens"],
            ["PPI", "Bora Embalagens"],
            ["04", "ABN Embalagens"],
            ["GLR", "ABN Embalagens"],
        ],
    )
    add_bullets(
        document,
        [
            "Se o armazém estiver selecionado, a empresa é preenchida automaticamente.",
            "Se o armazém não for informado, o lote entra como pendente.",
            "O armazém pode ser preenchido ou corrigido depois, inclusive em lote.",
        ],
    )

    document.add_heading("5.4 Tratamento de duplicidade", level=2)
    add_paragraph(
        document,
        "Quando o mesmo lote já existe na lista e ele é lido novamente, o DelCod não bloqueia de forma silenciosa. O sistema exibe o aviso 'Código duplicado' e pergunta se o usuário deseja continuar assim mesmo.",
    )
    add_bullets(
        document,
        [
            "Cancelar: interrompe o lançamento duplicado.",
            "Continuar: permite gravar o lote novamente mesmo já existindo na lista.",
        ],
    )

    document.add_heading("6. Lista global", level=1)
    add_paragraph(
        document,
        "A lista global concentra todos os lotes de bobina registrados. Cada item mostra o lote, o armazém, a empresa derivada, a origem do lançamento e o status operacional.",
    )
    add_bullets(
        document,
        [
            "Origem: câmera, importação ou manual.",
            "Armazém: mostra o código alocado ou 'Não informado'.",
            "Empresa: mostra Bora Embalagens, ABN Embalagens ou 'Pendente'.",
            "Status: mostra 'Completo', 'Sem armazém alocado' ou 'Armazém não mapeado'.",
            "Ações por item: editar e excluir.",
        ],
    )

    document.add_heading("6.1 Editar lote de bobina", level=2)
    add_numbered_steps(
        document,
        [
            "Na lista global, localize o lote desejado.",
            "Clique ou toque no ícone 'Editar'.",
            "Altere o valor do lote, se necessário.",
            "Selecione um novo armazém, mantenha o atual ou remova o armazém.",
            "Clique em 'Salvar'.",
        ],
    )
    add_paragraph(
        document,
        "Ao remover o armazém de um item já alocado, o lote volta para o estado pendente.",
    )

    document.add_heading("6.2 Excluir lote", level=2)
    add_numbered_steps(
        document,
        [
            "Localize o item na lista global.",
            "Clique ou toque no ícone 'Excluir'.",
            "Confirme a remoção quando desejar excluir o lote da lista.",
        ],
    )

    document.add_heading("6.3 Seleção múltipla", level=2)
    add_paragraph(
        document,
        "Cada linha da lista possui uma caixa de seleção. Esses marcadores permitem exportar apenas parte da base ou aplicar um armazém em vários lotes de uma vez.",
    )
    add_bullets(
        document,
        [
            "Selecionar individualmente: marque os itens desejados um a um.",
            "Selecionar todos: use o botão da área de ações para marcar toda a lista atual.",
            "Limpar seleção: o mesmo botão alterna para remover a seleção total.",
        ],
    )

    document.add_heading("6.4 Alocar armazém em lote", level=2)
    add_paragraph(
        document,
        "Quando vários lotes precisam receber o mesmo armazém, não é necessário editar um por um. O DelCod permite aplicar a alocação em lote a partir da seleção múltipla.",
    )
    add_numbered_steps(
        document,
        [
            "Selecione os lotes desejados na lista global.",
            "Clique em 'Alocar armazém'.",
            "Escolha o armazém que será aplicado aos itens selecionados.",
            "Se existirem lotes que já têm armazém, o sistema mostrará um resumo do que será reescrito.",
            "Escolha entre 'Somente pendentes' ou 'Reescrever selecionados'.",
        ],
    )
    add_bullets(
        document,
        [
            "Somente pendentes: aplica o novo armazém apenas aos lotes que ainda não tinham alocação.",
            "Reescrever selecionados: substitui o armazém atual dos itens destacados no resumo.",
        ],
    )

    document.add_heading("7. Importação de arquivos", level=1)
    add_paragraph(
        document,
        "O DelCod importa arquivos nos formatos Excel (.xlsx) e CSV. A importação foi preparada para funcionar tanto no navegador quanto no Android.",
    )
    add_bullets(
        document,
        [
            "O arquivo pode ter cabeçalho ou não.",
            "O sistema tenta reconhecer automaticamente as colunas de 'Lote de Bobina' e 'Armazém'.",
            "Apenas essas duas informações são utilizadas; as demais colunas são ignoradas.",
            "Se a coluna de armazém estiver presente, a alocação é aplicada automaticamente durante a importação.",
        ],
    )
    document.add_heading("7.1 Passo a passo da importação", level=2)
    add_numbered_steps(
        document,
        [
            "Clique em 'Importar arquivo'.",
            "Escolha um arquivo .xlsx ou .csv.",
            "Confirme se a primeira linha é cabeçalho.",
            "Verifique a coluna de lotes de bobina.",
            "Confirme ou ajuste a coluna de armazém, quando existir.",
            "Revise a prévia e o resumo do lote.",
            "Escolha 'Importar somente os novos' ou 'Importar tudo mesmo assim'.",
        ],
    )
    add_paragraph(
        document,
        "No resumo da importação, o sistema informa total lido, novos, duplicados, coluna usada e coluna de armazém. Se o mesmo lote aparecer repetido dentro do próprio arquivo, somente a primeira ocorrência entra como nova; as demais contam como duplicadas.",
    )

    document.add_heading("8. Exportação de dados", level=1)
    add_paragraph(
        document,
        "A exportação funciona com toda a lista ou apenas com os itens selecionados. Se nada estiver selecionado, o DelCod exporta todos os lotes ativos.",
    )
    document.add_heading("8.1 Exportar Excel", level=2)
    add_bullets(
        document,
        [
            "Botão: 'Exportar Excel'.",
            "Formato do arquivo: .xlsx.",
            "Nome da aba do arquivo: Leituras.",
            "Colunas exportadas: Lote de Bobina, Armazém, Empresa e Status.",
        ],
    )
    document.add_heading("8.2 Exportar PDF", level=2)
    add_bullets(
        document,
        [
            "Botão: 'Exportar PDF'.",
            "Formato do arquivo: .pdf.",
            "Conteúdo: tabela com Lote de Bobina, Armazém, Empresa e Status.",
            "Uso típico: conferência rápida, impressão e envio interno.",
        ],
    )

    document.add_heading("9. Limpar tudo", level=1)
    add_paragraph(
        document,
        "O botão 'Limpar tudo' remove todos os lotes ativos da lista global. Essa ação pede confirmação antes de executar.",
    )
    add_bullets(
        document,
        [
            "Use essa função apenas quando realmente desejar reiniciar a lista operacional.",
            "A ação é global e afeta todos os itens ativos visíveis no sistema.",
            "O sistema mostra um diálogo de confirmação antes da limpeza.",
        ],
    )

    document.add_heading("10. Funcionamento offline e sincronização", level=1)
    add_paragraph(
        document,
        "O DelCod foi construído com operação offline. Isso significa que o usuário pode continuar lançando lotes mesmo sem internet. Quando a conexão volta, o sistema tenta enviar automaticamente as operações pendentes para o banco de dados e atualizar a lista global.",
    )
    add_bullets(
        document,
        [
            "Sem internet: o aplicativo continua registrando localmente.",
            "Com internet: as alterações são sincronizadas e refletidas nos demais dispositivos.",
            "Banner de status: sempre indica se o sistema está sincronizado, sincronizando, offline ou com falha.",
            "Pendências: o banner informa quantas operações ainda aguardam envio.",
        ],
    )

    document.add_heading("11. Regras importantes do processo", level=1)
    add_bullets(
        document,
        [
            "Lote de Bobina é o valor principal do registro e aparece como referência em toda a operação.",
            "Armazém pode ser deixado em branco para não travar a leitura, mas o lote ficará pendente.",
            "Empresa é sempre derivada a partir do armazém, não sendo digitada manualmente.",
            "Lotes duplicados podem ser permitidos, desde que o usuário confirme.",
            "Importações grandes devem ser revisadas pelo resumo do lote antes da confirmação final.",
        ],
    )

    document.add_heading("12. Boas práticas de uso", level=1)
    add_bullets(
        document,
        [
            "No Android, mantenha o código de barras centralizado e com boa iluminação para leitura mais rápida.",
            "Quando souber o armazém antes da leitura, selecione-o no painel para economizar retrabalho.",
            "Use a seleção múltipla para alocação em lote sempre que vários itens forem do mesmo armazém.",
            "Revise os pendentes regularmente para evitar itens sem empresa derivada.",
            "Em importações, confira se o cabeçalho e a coluna do armazém foram reconhecidos corretamente.",
        ],
    )

    document.add_heading("13. Dúvidas e situações comuns", level=1)
    add_simple_table(
        document,
        ["Situação", "O que fazer"],
        [
            ["O lote entrou sem armazém.", "Edite o item individualmente ou use seleção múltipla e 'Alocar armazém'."],
            ["A leitura marcou duplicidade.", "Revise se o lote já existe. Se a repetição for intencional, confirme a continuação."],
            ["A planilha tem várias colunas.", "Na importação, selecione apenas a coluna do lote e, se existir, a coluna do armazém."],
            ["A internet caiu.", "Continue operando normalmente. O banner mostrará Offline e as mudanças serão sincronizadas depois."],
            ["Preciso exportar só parte da lista.", "Selecione os itens desejados e use 'Exportar Excel' ou 'Exportar PDF'."],
        ],
    )

    document.add_heading("14. Observação sobre materiais visuais", level=1)
    add_paragraph(
        document,
        "Esta versão do manual já inclui imagens reais do produto. Caso seja necessário, uma próxima revisão pode adicionar vídeos curtos de treinamento ou links/QR Codes para vídeos de operação, principalmente para leitura por câmera, importação e alocação em lote.",
    )

    return document


def main():
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(TARGET)
    print(f"Documento gerado em: {TARGET}")


if __name__ == "__main__":
    main()
