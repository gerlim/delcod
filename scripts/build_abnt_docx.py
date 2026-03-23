from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Projeto Leitor de cod de barras")
SOURCE = ROOT / "docs" / "superpowers" / "specs" / "2026-03-23-barcode-app-design.md"
TARGET = ROOT / "docs" / "superpowers" / "specs" / "2026-03-23-barcode-app-design-abnt.docx"


def set_run_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)

    set_run_font(run, size=10)


def restart_page_numbering(section, start: int = 1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def ensure_style(document: Document, name: str):
    if name in document.styles:
        return document.styles[name]
    return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    metadata = ensure_style(document, "Metadata")
    metadata.base_style = normal
    metadata.font.name = "Times New Roman"
    metadata._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    metadata.font.size = Pt(12)
    metadata.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    metadata.paragraph_format.line_spacing = 1.5
    metadata.paragraph_format.first_line_indent = Cm(0)
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(0)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(12)
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.line_spacing = 1.5
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.line_spacing = 1.5
    heading2.paragraph_format.first_line_indent = Cm(0)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.italic = True
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.line_spacing = 1.5
    heading3.paragraph_format.first_line_indent = Cm(0)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

    bullet = ensure_style(document, "ABNTBullet")
    bullet.base_style = normal
    bullet.font.name = "Times New Roman"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    bullet.font.size = Pt(12)
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullet.paragraph_format.line_spacing = 1.5
    bullet.paragraph_format.left_indent = Cm(1.25)
    bullet.paragraph_format.first_line_indent = Cm(-0.63)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(0)

    numbered = ensure_style(document, "ABNTNumbered")
    numbered.base_style = normal
    numbered.font.name = "Times New Roman"
    numbered._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    numbered.font.size = Pt(12)
    numbered.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    numbered.paragraph_format.line_spacing = 1.5
    numbered.paragraph_format.left_indent = Cm(1.25)
    numbered.paragraph_format.first_line_indent = Cm(-0.85)
    numbered.paragraph_format.space_before = Pt(0)
    numbered.paragraph_format.space_after = Pt(0)

    cover_title = ensure_style(document, "ABNTCoverTitle")
    cover_title.font.name = "Times New Roman"
    cover_title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cover_title.font.size = Pt(16)
    cover_title.font.bold = True
    cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.line_spacing = 1.5
    cover_title.paragraph_format.first_line_indent = Cm(0)
    cover_title.paragraph_format.space_before = Pt(0)
    cover_title.paragraph_format.space_after = Pt(0)

    cover_subtitle = ensure_style(document, "ABNTCoverSubtitle")
    cover_subtitle.font.name = "Times New Roman"
    cover_subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cover_subtitle.font.size = Pt(12)
    cover_subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle.paragraph_format.line_spacing = 1.5
    cover_subtitle.paragraph_format.first_line_indent = Cm(0)
    cover_subtitle.paragraph_format.space_before = Pt(0)
    cover_subtitle.paragraph_format.space_after = Pt(0)


def clean_text(text: str) -> str:
    return text.replace("`", "").strip()


def add_cover(document: Document, title: str, date_text: str, status_text: str):
    for _ in range(8):
        document.add_paragraph("")

    p = document.add_paragraph(style="ABNTCoverTitle")
    p.add_run(title.upper())

    p = document.add_paragraph(style="ABNTCoverSubtitle")
    p.add_run("Especificacao tecnica e executiva")

    p = document.add_paragraph(style="ABNTCoverSubtitle")
    p.add_run(status_text)

    for _ in range(12):
        document.add_paragraph("")

    p = document.add_paragraph(style="ABNTCoverSubtitle")
    p.add_run(date_text)


def add_metadata_paragraph(document: Document, line: str):
    p = document.add_paragraph(style="Metadata")
    label, value = line.split(":", 1)
    label_run = p.add_run(f"{label.strip()}: ")
    set_run_font(label_run, bold=True)
    value_run = p.add_run(value.strip())
    set_run_font(value_run)


def flush_paragraph(document: Document, buffer: list[str]):
    if not buffer:
        return
    text = clean_text(" ".join(part.strip() for part in buffer))
    if text:
        p = document.add_paragraph(style="Normal")
        p.add_run(text)
    buffer.clear()


def add_bullet(document: Document, text: str):
    p = document.add_paragraph(style="ABNTBullet")
    run = p.add_run(f"• {clean_text(text)}")
    set_run_font(run)


def add_numbered(document: Document, text: str):
    p = document.add_paragraph(style="ABNTNumbered")
    run = p.add_run(clean_text(text))
    set_run_font(run)


def add_heading(document: Document, level: int, text: str):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
    p = document.add_paragraph(style=style)
    heading_text = clean_text(text)
    if level == 1:
        heading_text = heading_text.upper()
    p.add_run(heading_text)


def build_body(document: Document, lines: list[str]):
    buffer: list[str] = []

    heading_re = re.compile(r"^(#{2,4})\s+(.*)$")
    ordered_re = re.compile(r"^\d+\.\s+.+$")

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            flush_paragraph(document, buffer)
            continue

        heading_match = heading_re.match(line)
        if heading_match:
            flush_paragraph(document, buffer)
            hashes, text = heading_match.groups()
            add_heading(document, min(len(hashes) - 1, 3), text)
            continue

        if line.startswith("- "):
            flush_paragraph(document, buffer)
            add_bullet(document, line[2:])
            continue

        if ordered_re.match(line):
            flush_paragraph(document, buffer)
            add_numbered(document, line)
            continue

        if ":" in line and not buffer and re.match(r"^[A-Za-zÀ-ÿ0-9 ._-]+:\s+.+$", line):
            flush_paragraph(document, buffer)
            add_metadata_paragraph(document, line)
            continue

        buffer.append(line)

    flush_paragraph(document, buffer)


def main():
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    title = clean_text(lines[0].removeprefix("# ").strip())

    date_text = "23 de marco de 2026"
    status_text = "Design aprovado para planejamento"
    for line in lines[1:10]:
        if line.startswith("Data:"):
            raw = line.split(":", 1)[1].strip()
            if raw == "2026-03-23":
                date_text = "23 de marco de 2026"
        if line.startswith("Status:"):
            status_text = line.split(":", 1)[1].strip()

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    add_cover(document, title, date_text, status_text)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    body_section.header.is_linked_to_previous = False
    restart_page_numbering(body_section, start=1)
    add_page_number(body_section.header.paragraphs[0])

    body_lines = lines[1:]
    build_body(document, body_lines)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
